#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成气候风险压测系统需求说明书（Word）"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SHOT_DIR = ROOT / 'requirements-screenshots'
OUT_FILE = ROOT / '气候风险压测系统需求说明书.docx'

MENUS = [
    {
        'menu': '数据处理',
        'functions': [
            {
                'name': '新建任务',
                'desc': '用于创建新的数据处理任务，录入任务名称、报告年度、贷款类型、贷款地区、压测目的及涉及行业等基本信息，为后续同步信贷与财务数据建立压测样本范围。',
                'shot': '02-data-process-create.png',
            },
            {
                'name': '查询 / 重置',
                'desc': '支持按任务名称、报告年度、贷款类型、贷款地区筛选数据处理任务列表，便于管理员快速定位目标任务；重置可清空筛选条件恢复全量列表。',
                'shot': '01-data-process-list.png',
            },
            {
                'name': '查看 / 编辑 / 删除',
                'desc': '查看以只读方式浏览任务详情与同步结果；编辑可修改任务基本信息或继续处理数据；删除用于清理草稿或不再使用的任务记录。',
                'shot': '01-data-process-list.png',
            },
            {
                'name': '同步贷款数据',
                'desc': '从信贷系统按任务概览中的贷款地区、报告年度、涉及行业、贷款类型拉取高碳行业客户贷款样本，生成客户基础信息表初始数据。',
                'shot': '03-data-process-financial-sync.png',
            },
            {
                'name': '同步财务数据',
                'desc': '在贷款数据同步完成后，从财务系统补充客户收入、利润、温室气体排放等字段，支撑后续碳排放与财务传导计算。',
                'shot': '03-data-process-financial-sync.png',
            },
            {
                'name': '同步内部PD数据',
                'desc': '针对行内压测场景，从内部系统同步客户 PD、LGD 等风险参数，为后续 PD/LGD 计算步骤提供行内口径数据。',
                'shot': '03-data-process-financial-sync.png',
            },
            {
                'name': '状态筛选',
                'desc': '按可使用、需计算、无法处理、已排除等状态筛选客户样本，便于聚焦待处理或异常数据。',
                'shot': '04-data-process-customer-table.png',
            },
            {
                'name': '导出 / 导入客户基础信息',
                'desc': '导出当前客户基础信息表用于线下核对或补录；导入可将补录后的 Excel 结果回写系统，更新样本可用状态。',
                'shot': '04-data-process-customer-table.png',
            },
            {
                'name': '编辑温室气体排放数据',
                'desc': '对尚未完成的数据处理任务，支持手工维护或修正客户温室气体排放量，满足碳排放测算对排放因子的输入要求。',
                'shot': '04-data-process-customer-table.png',
            },
            {
                'name': '行业甄别确认',
                'desc': '当系统识别到行业映射歧义客户时，引导业务人员确认标准行业归属，确保后续因子匹配与压测口径一致。',
                'shot': '04-data-process-customer-table.png',
            },
            {
                'name': '一键处理',
                'desc': '将符合条件的「需计算」样本批量标记为可使用，减少逐条确认操作，加快数据处理进度。',
                'shot': '04-data-process-customer-table.png',
            },
            {
                'name': '导出 / 导入内部PD/LGD数据',
                'desc': '支持内部 PD/LGD 明细的导出与 Excel 导入补录，满足行内压测对风险参数维护的业务需要。',
                'shot': '03-data-process-financial-sync.png',
            },
            {
                'name': '导出参试银行基础信息',
                'desc': '将按当前客户清单动态汇总的参试银行基础信息表导出为文件，供监管报送或内部分析使用。',
                'shot': '05-data-process-bank-basic.png',
            },
            {
                'name': '导出资本与拨备监管指标 / 编辑数据',
                'desc': '展示并导出核心资本、拨备率等监管指标；其中拨备相关监管要求支持手工调整，以反映参试银行实际监管口径。',
                'shot': '05-data-process-bank-basic.png',
            },
            {
                'name': '操作日志',
                'desc': '记录任务创建、同步、导入导出等关键操作，便于审计追溯与问题排查。',
                'shot': '03-data-process-financial-sync.png',
            },
        ],
    },
    {
        'menu': '情景分析',
        'functions': [
            {
                'name': '新建压测任务',
                'desc': '基于已完成的数据处理任务或导入 Excel 创建压测任务，作为情景分析、财务传导及后续压测步骤的统一载体。',
                'shot': '08-scenario-analysis-list.png',
            },
            {
                'name': '查询 / 重置',
                'desc': '按压测任务名称等条件检索任务列表，支持管理员在大量压测任务中快速定位目标。',
                'shot': '08-scenario-analysis-list.png',
            },
            {
                'name': '进入压测 / 查看',
                'desc': '进入指定压测任务的流水线步骤页，继续未完成压测或查看已完成任务的情景参数与中间结果。',
                'shot': '08-scenario-analysis-list.png',
            },
            {
                'name': '情景多选',
                'desc': '在同一压测任务内勾选基准、温室世界、有序转型等情景，满足监管要求的多情景并行压测需求。',
                'shot': '09-scenario-analysis-params.png',
            },
            {
                'name': '情景参数录入',
                'desc': '为各选中情景录入起始/结束年份、行业增长率、免费配额、碳价、资产负债率等参数，右侧实时预览碳排放费用走势。',
                'shot': '09-scenario-analysis-params.png',
            },
            {
                'name': '执行压测',
                'desc': '依据已录入情景参数对客户样本执行碳费用传导试算，生成情景分析阶段的初步压测结果。',
                'shot': '09-scenario-analysis-params.png',
            },
            {
                'name': '下一步：财务传导',
                'desc': '情景分析完成后进入财务传导步骤，将碳费用影响进一步传导至企业财务指标与银行资产质量。',
                'shot': '09-scenario-analysis-params.png',
            },
        ],
    },
    {
        'menu': '财务传导',
        'functions': [
            {
                'name': '新建压测任务',
                'desc': '与情景分析共用压测任务列表，可在财务传导菜单下新建或进入任务，按流水线推进压测。',
                'shot': '10-stress-fin-trans.png',
            },
            {
                'name': '执行财务传导',
                'desc': '将情景分析得到的碳费用结果传导至客户收入、成本、净利润及资产负债率等财务指标，形成财务传导中间表。',
                'shot': '10-stress-fin-trans.png',
            },
            {
                'name': '下一步：PD/LGD计算',
                'desc': '财务传导完成后进入 PD/LGD 计算步骤，为违约概率与损失率调整提供输入。',
                'shot': '10-stress-fin-trans.png',
            },
        ],
    },
    {
        'menu': 'PD/LGD计算',
        'functions': [
            {
                'name': '执行 PD/LGD 计算',
                'desc': '结合财务传导结果与内部/监管 PD、LGD 参数，计算压测后客户违约概率与损失率变化，支撑资产质量评估。',
                'shot': '11-stress-pd-lgd.png',
            },
            {
                'name': '下一步：不良和拨备计算',
                'desc': 'PD/LGD 计算完成后进入不良和拨备计算步骤，汇总不良贷款与拨备计提影响。',
                'shot': '11-stress-pd-lgd.png',
            },
        ],
    },
    {
        'menu': '不良和拨备计算',
        'functions': [
            {
                'name': '执行不良和拨备计算',
                'desc': '基于前述步骤结果计算压测情景下不良贷款余额、拨备计提金额等监管关注指标，完成压测流水线末端计算。',
                'shot': '12-stress-npl-prov.png',
            },
            {
                'name': '查看压测结果分析',
                'desc': '压测完成后跳转至结果分析模块，查看组合层面与行业层面的汇总指标及违约客户监控。',
                'shot': '12-stress-npl-prov.png',
            },
        ],
    },
    {
        'menu': '压测结果分析',
        'functions': [
            {
                'name': '结果来源 / 年份 / 情景筛选',
                'desc': '从已完成的数据处理任务或压测任务中选择结果集，并按年份、情景维度切换分析视角。',
                'shot': '13-results-analysis.png',
            },
            {
                'name': '导出明细',
                'desc': '将当前筛选条件下的压测明细结果导出为 Excel，并自动写入导出记录留痕，满足监管报送与内部分发需要。',
                'shot': '13-results-analysis.png',
            },
            {
                'name': '应用报送',
                'desc': '对已完成监管口径数据处理任务，进入应用报送页面，生成监管 Excel 文件包或下发内部风险预警。',
                'shot': '13-results-analysis.png',
            },
            {
                'name': '监管指标汇总表',
                'desc': '按监管要求展示资本、拨备、不良等关键指标的压测前后对比汇总，支撑监管沟通与内部汇报。',
                'shot': '13-results-analysis.png',
            },
            {
                'name': '违约客户监控 / 一键下发预警',
                'desc': '按行业展示压测新增违约/不良客户分布，支持导出监控数据并一键下发风险预警至相关管理条线。',
                'shot': '13-results-analysis.png',
            },
        ],
    },
    {
        'menu': '导出记录',
        'functions': [
            {
                'name': '查询 / 重置',
                'desc': '按导出类型、操作人、时间范围检索历史导出记录，便于审计与重复下载。',
                'shot': '14-exports.png',
            },
            {
                'name': '下载',
                'desc': '对历史导出文件进行再次下载，避免重复执行压测或明细导出操作。',
                'shot': '14-exports.png',
            },
        ],
    },
    {
        'menu': '基础配置 — 因子库管理',
        'functions': [
            {
                'name': '新增因子',
                'desc': '维护各行业碳排放因子编码、国标代码、因子值及生效版本，为碳排放测算提供基础参数库。',
                'shot': '16-factor-modal.png',
            },
            {
                'name': '查看 / 编辑 / 启用 / 停用 / 删除',
                'desc': '对因子进行全生命周期管理，确保仅生效因子参与压测计算，停用或删除过期因子以控制口径风险。',
                'shot': '15-factors.png',
            },
        ],
    },
    {
        'menu': '基础配置 — 行业映射关系',
        'functions': [
            {
                'name': '新增映射',
                'desc': '配置信贷系统接口行业与压测标准行业、国标代码的对应关系，解决样本行业口径不一致问题。',
                'shot': '17-mappings.png',
            },
            {
                'name': '查看 / 编辑 / 启用 / 停用 / 删除',
                'desc': '维护映射规则版本与状态，保障数据处理同步阶段能准确识别客户所属标准行业。',
                'shot': '17-mappings.png',
            },
        ],
    },
    {
        'menu': '基础配置 — 机场吞吐量维护',
        'functions': [
            {
                'name': '保存记录 / 编辑 / 删除',
                'desc': '维护机场企业旅客与货邮吞吐量数据，供机场类客户碳排放测算时自动调取，减少手工补录。',
                'shot': '18-airport-throughput.png',
            },
        ],
    },
    {
        'menu': '基础配置 — 计算方法说明',
        'functions': [
            {
                'name': '计算方法文档查阅',
                'desc': '集中展示碳排放、财务传导、PD/LGD、不良拨备等核心计算公式与口径说明，便于业务人员理解压测逻辑。',
                'shot': '19-calc-doc.png',
            },
        ],
    },
    {
        'menu': '菜单权限',
        'functions': [
            {
                'name': '菜单显示开关',
                'desc': '按角色或部署需要控制各一级/二级菜单的可见性，实现不同用户群体看到适配的功能范围。',
                'shot': '20-menu-perms.png',
            },
        ],
    },
]


def set_run_font(run, size=11, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x14, 0x25, 0x28)


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, 16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size, bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def add_image(doc, shot_name, width=5.8):
    path = SHOT_DIR / shot_name
    if not path.exists():
        add_para(doc, f'（截图缺失：{shot_name}）', size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'图：{shot_name.replace(".png", "").replace("-", " ")}')
    set_run_font(r, 9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69)
    sec.page_width = Inches(8.27)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('气候风险压测系统\n需求说明书')
    set_run_font(r, 22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('华夏银行绿金系统 · 原型 v1.0')
    set_run_font(r2, 12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = date_p.add_run('2026年7月')
    set_run_font(r3, 12)

    doc.add_page_break()

    # 目录说明
    add_title(doc, '1. 文档说明', 1)
    add_para(doc, '本文档基于气候风险压测系统（完整版原型 v1.0）梳理现有菜单与功能操作，描述各功能点的业务目的与实现需求，并附系统最新界面截图，供需求评审、验收测试及后续开发参考。')

    add_title(doc, '2. 功能清单', 1)
    add_para(doc, '下表汇总系统一级菜单、功能操作及简要说明。')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['序号', '系统菜单', '功能名称', '功能说明']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, 10, bold=True)

    idx = 1
    for menu_block in MENUS:
        menu_name = menu_block['menu']
        for fn in menu_block['functions']:
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = menu_name
            row[2].text = fn['name']
            row[3].text = fn['desc']
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, 10)
            idx += 1

    doc.add_page_break()

    # 详细章节
    add_title(doc, '3. 功能详细说明与界面截图', 1)
    sec_no = 1
    for menu_block in MENUS:
        add_title(doc, f'3.{sec_no} {menu_block["menu"]}', 2)
        sec_no += 1
        for i, fn in enumerate(menu_block['functions'], 1):
            add_title(doc, f'{fn["name"]}', 3)
            add_para(doc, fn['desc'])
            add_image(doc, fn['shot'])
            doc.add_paragraph()

    doc.save(OUT_FILE)
    print(f'Generated: {OUT_FILE}')


if __name__ == '__main__':
    build()
